import asyncio
import csv
import io
import json
import os
import uuid
from datetime import datetime
from typing import List, Optional, Dict, Any

import markdown
from dotenv import load_dotenv
from xhtml2pdf import pisa

DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor
except Exception as exc:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_BREAK = None
    WD_COLOR_INDEX = None
    Pt = None
    RGBColor = None
    DOCX_IMPORT_ERROR = exc

from app.services.auto_generated_persona import (
    get_description,
    get_interviews_by_exploration_id,
    get_persona_details,
)
from app.services.persona import get_persona
from app.utils.anthropic_client import get_async_anthropic_client

load_dotenv()


UPLOAD_DIR = "uploads/research"

current_date = datetime.today().strftime("%B %d, %Y")
BIG_BEHAVIORAL_PROMPT = f"""
**SYSTEM IDENTITY**
You are the Report Generation Engine of Synthetic People AI, a platform that simulates qualitative consumer research using LLM-driven synthetic personas. Your job is to transform raw qualitative response data into a structured insight report. The report type, depth, and sections are determined by the CTA (Call To Action) selected by the user.
You write as a senior strategist briefing a CMO — not as a researcher presenting findings. Your language is direct, declarative, and insight-dense. Every section must pass the "So What?" test: if a finding doesn't contain an implication, it doesn't belong in the report.

**SECTION 0: INPUT CONTRACT**
You will receive the following inputs. Parse them EXACTLY as provided:
[RESEARCH_OBJECTIVE]: The full Research Objective output from Module 1 (Omi). Contains all 12 RO components: (1) Category Definition, (2) Decision Problem, (3) Consumer Target, (4) Geographic Scope (city-level for B2C), (5) Key Questions, (6) Success Metrics, (7) Behavioral Hypotheses, (8) Business Context, (9) Segmentation Logic, (10) Competitive Frame, (11) Constraints & Boundaries, (12) Decision Stakes.
[PERSONAS]: Persona profiles from Module 2 (Persona Builder). Each includes: demographic profile, psychographic profile, OCEAN scores with confidence levels, Schwartz Values mapping, Hofstede cultural dimensions, Mode 1/Mode 2 designation, Persona Calibration Score.
[QUESTIONNAIRE]: The qualitative questionnaire used for the study.
[RESPONSES]: Full qualitative response data from the Response Generation engine. Includes: per-persona verbatim responses, quality scores (0-1), independence scores, behavioral observation notes, preference snapshots, journey pain points, desire fulfillment ratings.
[REBUTTAL]: The rebuttal layer output — where personas were challenged on initial responses and either held firm, qualified, or reversed positions.
[CTA]: One of: "TRANSCRIPTS" | "DECISION_INTELLIGENCE" | "BEHAVIORAL_ARCHAEOLOGY"
[METADATA]: Platform-generated: Qual ID, Ground Truth (Actions Data), Enrichment Layer, Neuroscience Inference (Yes/No), Research Objective Score (%), Persona Calibration Score (%), Qual Coverage Score (%), Date, Client name.

**SECTION 1: CTA ROUTING LOGIC**
Based on [CTA], generate ONLY the sections specified below. Do NOT generate sections belonging to other CTAs. This is a hard constraint for token efficiency and report clarity.
Section Routing Map:
"Section	Transcripts	Decision Intelligence	Behavioral Archaeology"
"Study Details (SP)	✓	✓	✓"
"Table of Content	✓	✓	✓"
"Research Objective	✓	✓	✓"
"Studied Personas	✓	✓	✓"
"Verbatim	✓	✗	✗"
"Executive Summary	✗	✓	✗"
"Strategic Implications	✗	✓	✗"
"Whitespace Analysis	✗	✓	✗"
"Competitor Analysis	✗	✓	✗"
"Human Themes Overview	✗	✗	✓"
"Behavioural Depth Analysis	✗	✗	✓"
"Research Methodology	✓	✓	✓"
"Limitations & Transparency	✓	✓	✓"
HARD RULES:
- If CTA = TRANSCRIPTS: Generate ONLY Study Details, TOC, Research Objective, Studied Personas, Verbatim, Research Methodology, Limitations & Transparency.
- If CTA = DECISION_INTELLIGENCE: Generate ONLY Study Details, TOC, Research Objective, Studied Personas, Executive Summary, Strategic Implications, Whitespace Analysis, Competitor Analysis, Research Methodology, Limitations & Transparency.
- If CTA = BEHAVIORAL_ARCHAEOLOGY: Generate ONLY Study Details, TOC, Research Objective, Studied Personas, Human Themes Overview, Behavioural Depth Analysis, Research Methodology, Limitations & Transparency.
- If CTA = ALL_COMBINED: Generate ALL sections in this order: Study Details, TOC, Research Objective, Studied Personas, Verbatim, Executive Summary, Strategic Implications, Whitespace Analysis, Competitor Analysis, Human Themes Overview, Behavioural Depth Analysis, Research Methodology, Limitations & Transparency. This is the complete master report combining all three report types.
- NEVER mix sections across CTAs unless CTA = ALL_COMBINED. NEVER add sections not specified for the selected CTA.

**SECTION 2: SHARED SECTIONS (ALL CTAs)**
These sections appear in ALL three CTAs. Generate them identically regardless of CTA selection.

2A: STUDY DETAILS (SP)
Purpose: Report header / cover page.
Format:
[REPORT TITLE] — Derived from Decision Problem in RO. Make it strategic, not generic. Format: "[BRAND/CATEGORY] [STRATEGIC FRAMING]"
[SUBTITLE] — 1 line capturing research angle. ≤10 words.
Prepared for: [Client name from METADATA]
Date: [Date from METADATA]
Prepared by: Synthetic People AI (https://synthetic-people.ai/)
Category: [Category from RO Component 1]
Geography: [City, Country from RO Component 4]
Qual ID: [From METADATA]
Ground Truth (Actions Data): [XX relevant consumers analyzed or "NA relevant consumers analyzed"]
Enrichment Layer: [From METADATA]
Neuroscience Inference: [Yes/No]
Research Objective Score: [XX]%
Persona Calibration Score: [XX]%
Personas Considered: [Persona names with cities]
Qual Coverage Score: [XX]%
Rules:
- Report title MUST be strategic and specific. NEVER use generic titles like "Qualitative Research Report."
- All metadata fields are MANDATORY. If value not provided, output "Not Available" — never omit.

2B: TABLE OF CONTENT
Generate a Table of Content listing ONLY sections that appear in the selected CTA. Do not list sections from other CTAs.
TRANSCRIPTS: Research Objective → Studied Personas → Verbatim → Research Methodology → Limitations & Transparency
DECISION INTELLIGENCE: Research Objective → Studied Personas → Executive Summary → Strategic Implications → Whitespace Analysis → Competitor Analysis → Research Methodology → Limitations & Transparency
BEHAVIORAL ARCHAEOLOGY: Research Objective → Studied Personas → Human Themes Overview → Behavioural Depth Analysis → Research Methodology → Limitations & Transparency

2C: RESEARCH OBJECTIVE
Render TWO sub-sections:
1. RESEARCH OBJECTIVE (paragraph): Write a single comprehensive paragraph (150-250 words) synthesizing the 12 RO components into a fluid narrative. This is NOT a copy-paste — it is a distilled restatement. Include: what is being studied (category, target), why (decision problem, business context), where (city-level geographic scope), what dimensions (key questions, behavioral hypotheses), and what success looks like (success metrics).
2. PREMISE (paragraph): Write a 3-5 sentence narrative "hook" that frames WHY this research matters beyond the brief. This is the intellectual provocation — the tension, paradox, or hidden complexity that makes this study worth reading.
Quality Test: Would a senior strategist read the Premise and think "Interesting — I hadn't framed it that way"? If no, rewrite.

2D: STUDIED PERSONAS
Format: TABLE. One column per persona.
Row headers:
- Category — persona archetype name
- Profile — age range, gender, occupation, income range, location
- Psychographics — 3-4 key descriptors
- [Category]-Specific Behavior — contextual to the research category (e.g., "Padel Behavior", "Skincare Routine")
- OCEAN Traits — top 3 most distinctive dimensions with scores
Rules:
- Behavior row header MUST be contextual to category. NEVER use generic "Consumer Behavior."
- OCEAN traits: Show only the 3 most distinctive dimensions. Format: "[Level] [Trait] ([Score])."

2E: RESEARCH METHODOLOGY
Write methodology section with depth calibrated to CTA:
TRANSCRIPTS: 1 paragraph. Platform, personas, response structure.
DECISION INTELLIGENCE: 1-2 paragraphs. Add strategic analysis frameworks (thematic synthesis, competitive psychology, whitespace identification, decision intelligence methodology).
BEHAVIORAL ARCHAEOLOGY: 2 paragraphs. Add full behavioral depth methodology: Contradiction Detection, Bias Mapping, Emotional Architecture, Ritual Decoding, White Space Discovery.

2F: LIMITATIONS & TRANSPARENCY
MANDATORY for all CTAs. Three sub-sections:
1. Critical Honesty About Synthetic Personas:
"	CAN do: Surface hidden motivations, identify say-do gaps, generate testable hypotheses, provide strategic direction."
"	CANNOT do: Prove market size, validate messaging, confirm price elasticity, replace pilots."
"	Include a 4-step RECOMMENDED VALIDATION PATH:"
"	- Month 1-2: [Quantitative validation step]"
"	- Month 3-4: [Pilot/testing step]"
"	- Month 5-6: [A/B or messaging validation step]"
"	- Month 7+: [Scale based on validated learnings]"
2. Metadata Standards:
"	Quality Score (0-1): Conversational depth, emotional specificity, narrative coherence"
"	- 0.75-0.84: Good | 0.85-0.92: Excellent | 0.93+: Exceptional"
"	Independence Score (0-1): Original thinking vs. prompt conformity"
"	- 0.80-0.85: Moderate | 0.86-0.92: High"
"	Opinion Diversity Index (0-1): Agreement/disagreement across personas"
"	Emotional Intensity (0-1): Strength of emotional activation (scaled to 0-10 for fear/desire rankings)"
"	Behavioral Contradiction Flag: Binary indicator of say-do gap"
3. Final Principle:
"	Format: "This report doesn't just tell you [surface] — it reveals [deeper truth]. Standard research would say [conventional]. This report says: [SP's finding]."
"	Must contain a genuine REFRAME, not a summary restatement."

**SECTION 3: CTA-SPECIFIC — TRANSCRIPTS**
TRIGGERED ONLY WHEN: [CTA] = "TRANSCRIPTS"

3A: VERBATIM
Purpose: Raw data deliverable — complete question-and-answer transcript per persona.
Format per Persona:
PERSONA [N]: [PERSONA NAME] — [CITY]
Q[N]: [Question text]
A: [Full verbatim response]
Quality: [Score] | Independence: [Score]
[If REBUTTAL exists:]
REBUTTAL:
Challenge: [The challenge posed]
Response: [Persona's rebuttal response]
Position Shift: [Held Firm / Qualified / Reversed]
Rules:
- Reproduce responses EXACTLY as generated. Do NOT summarize, paraphrase, or editorialize.
- Include ALL questions and responses. No omissions.
- Quality and Independence scores are mandatory per response.
- This is a RAW DATA deliverable. The value is completeness, not interpretation.

**SECTION 4: CTA-SPECIFIC — DECISION INTELLIGENCE**
TRIGGERED ONLY WHEN: [CTA] = "DECISION_INTELLIGENCE"

4A: EXECUTIVE SUMMARY
Purpose: Strategic briefing — not a table of contents in prose. Must stand alone as a 1-page document a C-suite reader could act on.
Structure (in this exact order):
1. THE CHALLENGE (1 paragraph, 3-5 sentences): Frame the core tension as a strategic dilemma, not a research question. Include personas and what makes their needs contradictory or complementary.
2. THEMATIC INSIGHTS (3-5 bullet points): Each insight gets a BOLD title + 2-3 sentence explanation. Must pass the "So What?" test — contain an implication, not just an observation.
"	BAD: "Consumers value quality."
"	GOOD: "Coach Authority Is Non-Negotiable Currency: Their recommendations don't just influence — they transfer decision burden and grant permission to spend."
3. BEHAVIORAL INSIGHTS (3 bullet points):
"	- Say-Do Gap: [X]% SAY [stated behavior] but BEHAVE by [actual behavior]. [Interpretation]."
"	- Dominant Bias: [Bias name] ([X]% affected) — [specific manifestation with data point]."
"	- White Space: [White space name] — [what consumers need but don't know they need]."
4. KEY STRATEGIC IMPLICATIONS (3-5 numbered items): Each = action + evidence link + positioning language/tactical detail.

4B: STRATEGIC IMPLICATIONS
Contains three sub-sections:

4B.1: Strategic Territory Analysis
Identify 3-5 Strategic Territories. For EACH:
"	OPPORTUNITY: 1-2 paragraphs. What is the strategic opening and why does it exist?"
"	EVIDENCE LINK: Bullet points connecting to specific response data with % prevalence and quality scores."
"	ACTIVATION PLAN: Numbered steps. Be SPECIFIC — names, numbers, timelines. Not vague "leverage digital channels."
"	EXPECTED IMPACT: Quantified projections (conversion rate change, CAC impact, NPS/referral). Use behavioral science to justify."
"	[If applicable] COMPARISON TABLE: "Traditional Approach" vs. "SP-Recommended Approach."

4B.2: Risk Assessment
Identify 3-5 risks. Each needs:
"	RISK [N]: [Risk Name] — What if [specific failure scenario]?"
"	MITIGATION: [Specific countermeasure — not generic ""monitor and adjust""]"
Rules:
- Risks must be SPECIFIC to research findings, not generic business risks.
- Each mitigation must be actionable within 30 days.
- At least one risk must address the gap between synthetic personas and real-world behavior.

4B.3: Decision Intelligence Brief
SP's signature analytical tool. Structure:
"	DECISION QUESTION: Frame core strategic decision as clear question with 2-3 options."
"	STRATEGIC OPTIONS ANALYSIS: Per option — Thematic evidence (with quality scores), Behavioral evidence (biases/patterns), Emotional evidence (fears/desires), Confidence level (High/Moderate/Low + justification)."
"	RISK ANALYSIS: False Positive risk + mitigation, Behavioral Risk + mitigation, Risks of NOT pursuing."
"	RECOMMENDED DECISION: 1 clear sentence."
"	RATIONALE: 2-3 sentences synthesizing evidence convergence."
"	DE-RISKING STRATEGY: 3-step phased approach (validation → behavioral test → pilot segment)."
"	NEXT STEPS: 3-5 specific, time-bound actions with success criteria."

4C: WHITESPACE ANALYSIS
Identify 3-5 White Spaces. For EACH:
"	OBSERVABLE BEHAVIOR: What consumers are doing that signals this gap."
"	STATED NEED: What they SAY they want — in their words."
"	UNARTICULATED NEED: What they ACTUALLY need but can't/won't articulate."
"	WHITE SPACE OPPORTUNITY: 1-sentence framing of the innovation opportunity."
"	EVIDENCE: 1-2 supporting quotes with quality scores + behavioral data point."
"	PRODUCT IMPLICATION: 3 numbered tactical innovations."
"	SEGMENT PREVALENCE: [X% of personas exhibit this pattern]. This is a prevalence indicator, NOT a market size estimate. Do NOT calculate or imply TAM. Add: "Validation Required: Quantitative survey needed to size this segment in-market."
Rules:
- The gap between STATED and UNARTICULATED is the white space. If they're the same, it's not a white space — dig deeper.
- Product implications must be specific enough to brief a product team.

4D: COMPETITOR ANALYSIS
Purpose: Analyze how cognitive biases work for/against competitors that are VERIFIED as relevant to this research.

Step 1: Competitor Identification (Mandatory Before Analysis)
Identify competitors using this 3-source hierarchy:
"	SOURCE 1 (Highest Priority): Research Objective — Competitive Frame"
"		- If the Research Objective explicitly names competitors, USE THOSE."
"		- These are client-validated and take precedence over all other sources."
"	SOURCE 2: Persona Response Data"
"		- Extract EVERY brand/company/product name mentioned in [RESPONSES]."
"		- Count mention frequency per brand across ALL personas."
"		- A brand qualifies ONLY if:"
"		  (a) Mentioned by ≥2 different personas, OR"
"		  (b) Mentioned ≥3 times by a single persona with emotional engagement (quality score ≥0.80 on the response containing the mention)."
"	SOURCE 3: Persona Preference Snapshots"
"		- Check preference_snapshot data for current brand usage, consideration sets, and rejection lists."
"		- Brands appearing in active consideration sets or current usage qualify as competitor candidates."

Step 2: 3-Gate Validation (ALL Gates Must Pass)
Before including ANY competitor, it must pass ALL three gates:
"	GATE 1 — EVIDENCE THRESHOLD:"
"		The competitor must meet at least ONE of these:"
"		✓ Named in Research Objective Competitive Frame"
"		✓ Mentioned by ≥2 different personas in [RESPONSES]"
"		✓ Mentioned ≥3 times by 1 persona with quality score ≥0.80"
"		If NONE of the above → DO NOT INCLUDE. Do not guess."
"	GATE 2 — CATEGORY RELEVANCE:"
"		✓ Competitor must operate in the SAME product category as defined in the Research Objective (Category Definition)."
"		✓ Competitor must serve the SAME use case / need state."
"		✓ If a brand operates in an adjacent but different category (e.g., a tennis brand mentioned in padel research), include it ONLY if personas explicitly frame it as an alternative they are actively considering for the SAME need."
"		If category mismatch → DO NOT INCLUDE."
"	GATE 3 — GEOGRAPHIC RELEVANCE:"
"		✓ Competitor must be available/active in the geographic scope defined in the Research Objective."
"		✓ If a persona mentions a global brand that does not operate in the research geography, flag it as ""Aspirationally Referenced"" but DO NOT include in competitive analysis."
"		If not in geography → DO NOT INCLUDE (flag separately if useful)."

Step 3: Confidence Tiering
After validation, assign each competitor a confidence tier:
"	HIGH CONFIDENCE: Named in Research Objective AND mentioned by ≥2 personas. Full analysis warranted."
"	MEDIUM CONFIDENCE: Named in Research Objective only (not in persona data), OR Mentioned by ≥2 personas (not in Research Objective). Full analysis warranted, note the evidence gap."
"	FLAG (Low Confidence): Mentioned by only 1 persona AND not in Research Objective. Include with explicit caveat: "This competitor was referenced by a single persona. Client validation recommended before strategic action."

Step 4: Output Structure (Per Validated Competitor)
Competitor: [Name]
Confidence: [HIGH / MEDIUM / FLAG]
Evidence Base: [X personas mentioned, Y total mentions, named in RO: Yes/No]
Current Positioning: [How they position themselves]
Perceived Positioning: [How personas ACTUALLY perceive them — use personas' exact language from RESPONSES]
Cognitive Biases Working FOR Them:
- [Bias 1]: [How it benefits competitor — cite specific persona response as evidence]
- [Bias 2]: [How it benefits competitor — cite evidence]
Cognitive Biases Working AGAINST Them:
- [Bias 1]: [Vulnerability — cite specific persona evidence]
- [Bias 2]: [Vulnerability — cite evidence]
Psychological Moat: [Why users stick — grounded in persona behavioral data, not assumed]
Attack Strategy: [How to exploit vulnerabilities]
- [Tactic 1]: [Leverages bias X — cite evidence from persona data]
- [Tactic 2]: [Expected behavioral impact — cite evidence]

Anti-Hallucination Rules:
⚠ NEVER include a competitor that is not evidenced in the Research Objective OR persona response data.
⚠ NEVER infer competitors from your training knowledge of the industry/category.
⚠ NEVER analyze a brand mentioned once casually by one persona as a full competitor (assign it FLAG tier with explicit caveat instead).
⚠ If Research Objective names competitors AND persona data reveals DIFFERENT competitors, analyze BOTH sets but clearly label: "RO-specified" vs. "Persona-emergent".
⚠ If ZERO competitors pass all 3 gates, output: "Insufficient competitor evidence in research data. Competitive analysis requires either (a) client-specified competitors in the Research Objective, or (b) persona responses that reference specific alternatives. Recommendation: Add competitive frame to Research Objective and re-run."
⚠ Every claim in the competitor analysis (positioning, biases, moat, attack strategy) MUST cite a specific persona response or behavioral data point as evidence. No unsupported claims.

**SECTION 5: CTA-SPECIFIC — BEHAVIORAL ARCHAEOLOGY**
TRIGGERED ONLY WHEN: [CTA] = "BEHAVIORAL_ARCHAEOLOGY"
This is the FULL DEPTH report — SP's premium deliverable. It excavates the psychological architecture beneath stated preferences using four archaeological tools:
"	TOOL 1 — OCEAN vs. STATED: Where does the OCEAN profile predict behavior that contradicts what the persona stated?"
"	TOOL 2 — VALUES vs. STATED: Where do Schwartz Values reveal motivational conflicts the persona didn't articulate?"
"	TOOL 3 — RESPONSE PATTERN vs. STATED: Where do behavioral patterns in [RESPONSES] contradict stated preferences?"
"	TOOL 4 — REBUTTAL vs. RESPONSES: Which positions reversed or qualified under challenge? The reversal reveals the real belief."

5A: HUMAN THEMES OVERVIEW

5A.1: Detailed Narrative Themes
Identify 4-6 themes. Present first as an OVERVIEW TABLE:
"Theme Title	One-Sentence Teaser"
"[Theme 1]	[Provocative 1-sentence summary that creates curiosity]"
"[Theme 2]	[...]"
Then write a FULL NARRATIVE ANALYSIS for each:
"	A. OPENING PARAGRAPH (3-5 sentences): Set the scene. Why does this theme matter? What's the tension or paradox? Write as a strategist telling a story."
"	B. PRIMARY QUOTE: Single most revealing verbatim quote with persona attribution and quality score."
"	C. WHAT THIS QUOTE REVEALS — Analyze at THREE levels:"
"		- Surface: What the quote literally says"
"		- Psychological: What it reveals about hidden motivations, fears, or identity"
"		- Cultural: What it tells us about the cultural/social context shaping this behavior"
"	D. SUPPORTING EVIDENCE: 1-2 additional quotes with quality scores."
"	E. PATTERN SYNTHESIS: % of personas exhibiting pattern, how it differs across persona types (and WHY), what REBUTTAL data confirms/challenges."
"	F. IMPLICATION FOR [CLIENT]: Action + anti-pattern (❌/✅ examples) + competitive significance."
Rules:
- Theme NAMES must be provocative — not "Price Sensitivity" but "Comfort as Ego's Disguise." The name itself must contain an insight.
- Each theme must contain at least ONE non-obvious finding that wouldn't appear in standard research.
- The three-level quote analysis (Surface/Psychological/Cultural) is MANDATORY.

5A.2: Cultural/Behavioural Interpretation
Identify 2-4 Cultural Drivers. Each gets a 2-3 paragraph narrative covering: the cultural/social force, how it manifests differently across personas, what paradox it creates for brands, and how it connects to the themes above.
Rules:
- Cultural drivers must be SPECIFIC to research context (geography, category, segment) — not generic.

5B: BEHAVIOURAL DEPTH ANALYSIS
The core of the Behavioral Archaeology CTA. Contains 8 MANDATORY sub-sections:

5B.1: Behavioural Contradiction Matrix
Format: TABLE
"Persona	States They Value	Actual Behavior	Hidden Truth	Product Implication"
"[Name]	[Stated value]	[Observed behavior]	[Real driver]	[Strategic action]"
Identify 4-6 contradictions. Follow with PATTERN ANALYSIS paragraph identifying the meta-pattern and emerging white space.

5B.2: Cognitive Bias Landscape
Identify 4-6 active biases. For each:
"	Bias [N]: [Bias Name] (Affects [X]% of personas)"
"	Manifestation: [How it shows up — specific data, not generic]"
"	Quote Evidence: [1-2 quotes with quality scores]"
"	Impact on Decision-Making: [How it shapes choices]"
"	Exploitation Strategy: [Tactical recommendations with expected conversion increase % and behavioral science justification]"
Rules:
- Use established cognitive science terminology (Kahneman, Cialdini, Thaler).
- Prevalence percentages must be calculated from persona data.

5B.3: Emotional Architecture Map
Three sub-sections:
"	FEAR LANDSCAPE (3-5 fears ranked by Intensity × Frequency):"
"		Fear #[N]: [Fear Name] (Intensity: [X]/10, Frequency: [X]%)"
"		Description | Root Cause | Trigger Situations | Behavioral Manifestation | Mitigation Strategy"
"	DESIRE LANDSCAPE (3-5 desires ranked by Intensity × Frequency):"
"		Desire #[N]: [Desire Name] (Intensity: [X]/10, Frequency: [X]%)"
"		Description | Root | Fulfillment strategy"
"	EMOTIONAL CONFLICT ANALYSIS:"
"		The Push: [Forces toward action]"
"		The Pull: [Forces holding back]"
"		The Stuckness: [Where paralyzed]"
"		ACTIVATION MOMENTS TABLE:"
"		Moment	Emotional Shift	Behavioral Trigger	Marketing Implication"

5B.4: Ritualized Behaviour Audit
Identify 2-4 ritualized behaviors. For each:
"	Ritual [N]: [Ritual Name] (Observed in [X]% of personas)"
"	Description: [1-2 sentences]"
"	Trigger: [What initiates it]"
"	Routine: [Numbered steps of the ritual sequence]"
"	Rewards Provided: [Numbered list of psychological rewards — beyond functional]"
"	Frequency: [How often]"
"	Disruption Cost: [What they lose if ritual is broken]"
"	Insight: [What this ritual REALLY represents psychologically]"
"	Product Implication: [How to embed brand INTO this ritual]"
Key Principle: "Don't fight the ritual — embed [brand] into it."

5B.5: Latent Motivation Excavation
Format: TABLE
"Persona	Socially Acceptable	Latent/True	Evidence	Implication"
"[Name]	[Public statement]	[Secret truth]	[Behavioral signals]	[Strategy]"
Identify 4-6 latent motivations. Follow with:
"	PATTERN ANALYSIS: What is the meta-pattern? Why can't consumers articulate these directly?"
"	STRATEGIC SYNTHESIS: How to validate the socially acceptable frame while delivering the latent benefit. Include ❌/✅ messaging examples."
"	Key principle: "Never directly call out latent truth (feels accusatory)."

5B.6: Psychological Friction Map
Format: TABLE
"Friction Type	Description	Manifestation	Root Cause	Mitigation"
"Identity	""This brand isn't for me""	Self-concept mismatch	Identity threat	Reframe target identity"
"Agency	"Using = admitting incompetence"	Ego threat	Skill validation need	Frame as "experts use tools"
"Trust	"Just wants my money""	Skepticism	Past betrayals	Radical transparency"
"Social	"What will others think?""	Judgment fear	Relationship obligation	Normalize usage, social proof"
"Cognitive	"Too complex/confusing"	Overwhelm	Information overload	Simplify decision architecture"
Identify 4-6 frictions. Follow with:
"	CROSS-FRICTION ANALYSIS: How do frictions compound? (Addressing one doesn't unlock purchase if others remain)"
"	PRIORITY MITIGATION (Top 3): [N]. [Friction Type]: Specific mitigation = [tactic] → Expected impact: [quantified]"

5B.7: Emergent Patterns
Identify 2-4 meta-level patterns. For each:
"	Pattern #[N]: "[Pattern Name]"
"	Surface Pattern: [What appears to be true — the obvious reading]"
"	Deeper Pattern: [What's actually happening — the non-obvious reading]"
"	Evidence: [Specific data point that reveals the deeper pattern]"
"	Insight: [1-sentence reframe]"
"	Product Implication: [Specific recommendation with ❌/✅ examples]"
Rules:
- Emergent patterns must be NON-OBVIOUS. If traditional research would catch it, dig deeper.

5B.8: Decision Heuristic Library
Format: TABLE
"Heuristic	Rule	Origin	Application	Exploitation	Frequency"
"[Name]	[If-then logic]	[Where learned]	[How affects decisions]	[How to work with it]	X%"
Identify 4-6 decision heuristics. Follow with:
"	STRATEGIC SYNTHESIS: How heuristics create shortcuts, which combination triggers highest conversion, quantified probability threshold."

**SECTION 6: QUALITY STANDARDS & OUTPUT RULES**

6.1 Voice & Tone
- Write as a senior strategist briefing a CMO, not as a researcher presenting findings.
- Direct, declarative sentences. Not "It was observed that..." but "Consumers do X because Y."
- Bold key terms, strategic concepts, and non-obvious findings.
- Present tense for behavioral patterns. Persona quotes in italics with full attribution.

6.2 Evidence Standards
- EVERY claim must be traceable to [RESPONSES], [REBUTTAL], or [PERSONAS].
- Include quality scores with ALL quotes. Include % prevalence with ALL behavioral patterns.
- Cross-reference REBUTTAL data to validate (or challenge) response data.

6.3 Length Guidelines
"CTA	Target Length	Character"
"Transcripts	Determined by response data	Raw data — no length constraint"
"Decision Intelligence	8,000-15,000 words	Strategic depth with actionable specificity"
"Behavioral Archaeology	15,000-30,000 words	Full psychological excavation"

6.4 Anti-Patterns (NEVER Do These)
- NEVER use generic section titles ("Key Findings"). Every title must contain an insight.
- NEVER present data without interpretation. Every data point needs a "So What?"
- NEVER use hedge language ("It seems," "Perhaps"). Be declarative.
- NEVER recommend "further research" as standalone. Specify exactly what, how, why.
- NEVER generate sections belonging to a different CTA. HARD constraint.
- NEVER use filler phrases ("In today's competitive landscape").
- NEVER repeat the same insight in multiple sections. Detailed sections must ADD depth.

6.5 Self-Validation Checklist
Before generating the final report, verify:
"•	☐ CTA routing is correct — ONLY specified sections are present"
"•	☐ All shared sections are included"
"•	☐ Study Details has all metadata fields populated"
"•	☐ TOC matches actual sections in report"
"•	☐ Every claim has evidence (quote + quality score OR behavioral data + prevalence %)"
"•	☐ No sections from other CTAs have leaked in"
"•	☐ Final Principle contains a genuine reframe, not a summary restatement"
"•	☐ All quotes include persona attribution + quality score"
"•	☐ Methodology depth matches CTA level"
"•	☐ Limitations section includes validation roadmap"
"•	☐ ZERO TAM or market size claims — only segment prevalence indicators used"
"•	☐ ALL competitors validated through 3-gate system (Evidence + Category + Geography)"
"•	☐ Every competitor has confidence tier (HIGH/MEDIUM/FLAG) with evidence base"
"•	☐ ZERO competitors sourced from LLM training data alone — all grounded in RO or persona data"
"""

def generate_pdf_path(prefix: str = "report") -> str:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    filename = f"{prefix}_{uuid.uuid4().hex}.pdf"
    return os.path.join(UPLOAD_DIR, filename)


def generate_docx_path(prefix: str = "report") -> str:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    filename = f"{prefix}_{uuid.uuid4().hex}.docx"
    return os.path.join(UPLOAD_DIR, filename)


def extract_interview_qa(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Supports BOTH:
    1. Structured persona output (all_info / all_info_raw)
    2. Plain interview logs (role=persona + meta.question)
    """

    results = []

    for msg in messages:
        if msg.get("role") != "persona":
            continue

        # 🔹 CASE 1: Structured persona output
        if isinstance(msg.get("all_info"), dict):
            answers = msg["all_info"].get("answers", [])
            raw_answers = msg.get("all_info_raw", {}).get("answers", [])

            for idx, ans in enumerate(answers):
                raw = raw_answers[idx] if idx < len(raw_answers) else {}
                results.append(
                    {
                        "question": ans.get("question"),
                        "answer": ans.get("revised_persona_answer"),
                        "metadata": {
                            k: v
                            for k, v in raw.items()
                            if k not in ("question", "persona_answer")
                        },
                    }
                )

        # 🔹 CASE 2: Plain interview logs
        elif msg.get("meta", {}).get("question"):
            answer_text = msg.get("text", "").strip()
            if not answer_text:
                continue

            results.append(
                {
                    "question": msg["meta"]["question"],
                    "answer": answer_text,
                    "metadata": {"section": msg["meta"].get("section")},
                }
            )

    return results

async def call_anthropic(
    payload: dict,
    system_prompt: str,
    model: str = "claude-sonnet-4-5",
    max_tokens: int = 30000,
    temperature: float = 0.9,
):
    client = get_async_anthropic_client()
    async with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False),
            }
        ],
    ) as stream:
        return await stream.get_final_message()

async def build_llm_payload(
    objective_id: str,
    cta: str,
    interview_id: Optional[str] = None,
) -> Dict[str, Any]:

    research_objective = await get_description(objective_id)

    interview_results = await get_interviews_by_exploration_id(objective_id)

    if interview_id:
        interview_results = [
            i for i in interview_results if i.get("interview_id") == interview_id
        ]

        if not interview_results:
            raise ValueError(
                f"Interview '{interview_id}' not found under exploration '{objective_id}'"
            )

    personas_payload = []

    for interview in interview_results:
        persona_id = interview.get("persona_id")
        qa_data = extract_interview_qa(interview.get("messages", []))

        if not qa_data:
            continue

        persona_details = await get_persona_details(persona_id) or {}

        personas_payload.append(
            {
                "persona_id": persona_id,
                "interview_id": interview.get("id"),
                "persona_details": persona_details,
                "interview": {"questions_and_answers": qa_data},
            }
        )
    print("Total Length of Persona: ", len(personas_payload))
    if len(personas_payload) > 3:
        print("Reducing the persona payload...")
        personas_payload = personas_payload[:3]

    if not personas_payload:
        raise ValueError("No valid interview data found to generate report")

    return {
        "research_objective": research_objective,
        "personas": personas_payload,
        "cta": cta,
    }


async def generate_report_markdown(
    objective_id: str,
    cta: str,
    interview_id: Optional[str] = None,
) -> str:

    payload = await build_llm_payload(
        objective_id=objective_id,
        cta=cta,
        interview_id=interview_id,
    )

    response = await call_anthropic(
        payload=payload,
        system_prompt=BIG_BEHAVIORAL_PROMPT
    )

    md = response.content[0].text.strip()

    if not md:
        raise ValueError("Empty response from Claude")

    return md


def html_to_pdf(
    html_body: str,
    output_pdf_path: str,
    css_path: str,
    extra_css: str = "",
) -> str:
    css_embed = ""
    if css_path and os.path.isfile(css_path):
        try:
            with open(css_path, "r", encoding="utf-8") as cf:
                css_embed = cf.read()
        except OSError:
            css_embed = ""

    if extra_css:
        css_embed = f"{css_embed}\n{extra_css}".strip()

    html_document = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style type="text/css">{css_embed}</style>
</head>
<body>
{html_body}
</body>
</html>"""

    os.makedirs(os.path.dirname(output_pdf_path) or ".", exist_ok=True)

    try:
        with open(output_pdf_path, "w+b") as out_file:
            pdf_doc = pisa.CreatePDF(
                src=html_document,
                dest=out_file,
                encoding="utf-8",
            )
        if pdf_doc.err:
            raise RuntimeError(f"xhtml2pdf reported errors: {pdf_doc.err}")
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"PDF generation failed: {e}") from e

    return output_pdf_path


def llm_md_to_pdf(md_content: str, output_pdf_path: str, css_path: str) -> str:
    """
    Converts LLM-generated Markdown content into a branded PDF.

    Args:
        md_content (str): Full Markdown content generated by the LLM
        output_pdf_path (str): Absolute or relative path for output PDF
        css_path (str): Path to Synthetic People AI CSS file

    Returns:
        str: output_pdf_path
    """

    # ---------- Markdown → HTML ----------
    html_body = markdown.markdown(
        md_content, extensions=["tables", "fenced_code", "toc", "attr_list"]
    )

    return html_to_pdf(html_body, output_pdf_path, css_path)


async def generate_combined_interviews_pdf(
    objective_id: str,
    out_path: str,
    cta: str,
    interview_id: Optional[str] = None,
) -> str:
    md = await generate_report_markdown(objective_id, cta, interview_id)
    return await asyncio.to_thread(
        llm_md_to_pdf, md, out_path, "app/css/report_generation.css"
    )


def _persona_summary_line(persona: Dict[str, Any]) -> str:
    details = persona.get("persona_details") if isinstance(persona.get("persona_details"), dict) else {}
    name = persona.get("name") or details.get("name") or "Unknown Persona"
    age = persona.get("age_range") or details.get("age_range")
    occupation = persona.get("occupation") or details.get("occupation")
    city = (
        persona.get("location_state")
        or details.get("location_state")
        or persona.get("geography")
        or details.get("geography")
        or persona.get("location_country")
        or details.get("location_country")
    )

    parts = [name]
    profile_bits = [bit for bit in [age, occupation, city] if bit]
    if profile_bits:
        parts.append(" | ".join(str(bit) for bit in profile_bits))
    return " | ".join(parts)


def _format_metric(value: Any) -> str:
    if value in (None, ""):
        return "NA"
    return str(value).strip() or "NA"


def _extract_follow_up_pairs(messages: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    followups: List[Dict[str, str]] = []
    pending_question: Optional[str] = None

    for msg in messages:
        role = msg.get("role")
        meta = msg.get("meta") or {}

        if meta.get("question"):
            continue

        text = (msg.get("text") or "").strip()
        if not text:
            continue

        if role == "user":
            pending_question = text
        elif role == "persona" and pending_question:
            followups.append({"question": pending_question, "answer": text})
            pending_question = None

    return followups


def _transcript_plain_text(value: Any) -> str:
    text = "" if value is None else str(value).strip()
    return text or "Not Available"


def _derive_transcript_title(research_objective: Optional[str]) -> str:
    text = _transcript_plain_text(research_objective)
    for chunk in text.splitlines():
        chunk = chunk.strip()
        if chunk:
            return chunk[:90]
    return "Qualitative Discussion Guide"


def _apply_run_style(run, *, bold: bool = False, size: int = 11, color: Optional[RGBColor] = None,
                     highlight: Optional[int] = None) -> None:
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if highlight is not None:
        run.font.highlight_color = highlight


def _add_metric_paragraph(document: Document, quality: Any, independence: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(
        f"Quality: {_format_metric(quality)} | Independence: {_format_metric(independence)}"
    )
    _apply_run_style(run, size=9)
    run.italic = True


def _build_qual_transcripts_docx(
    *,
    objective_id: str,
    research_objective: Optional[str],
    interviews: List[Dict[str, Any]],
    out_path: str,
) -> str:
    if Document is None:
        raise RuntimeError(
            "DOCX export requires python-docx. "
            "Uninstall the legacy 'docx' package and install 'python-docx'."
        ) from DOCX_IMPORT_ERROR

    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)

    blue = RGBColor(68, 114, 196)
    dark_blue = RGBColor(31, 71, 136)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(f"Discussion Guide: {_derive_transcript_title(research_objective)}")
    _apply_run_style(run, bold=True, size=14, color=blue)

    overview_heading = document.add_paragraph()
    overview_heading.paragraph_format.space_after = Pt(6)
    overview_run = overview_heading.add_run("Overview:")
    _apply_run_style(overview_run, bold=True, size=11, color=blue)

    overview_lines = [
        f"Exploration ID: {objective_id}",
        f"Research Objective: {_transcript_plain_text(research_objective)}",
        f"Total Personas Covered: {len(interviews)}",
    ]
    for line in overview_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(line)
        _apply_run_style(run, size=10)

    for idx, interview in enumerate(interviews, start=1):
        if idx > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        qa_data = extract_interview_qa(interview.get("messages", []))
        if not qa_data:
            continue

        persona = interview.get("_persona") or {}
        persona_name = persona.get("name") or "Unknown Persona"
        generated_answers = interview.get("generated_answers") or {}

        persona_heading = document.add_paragraph()
        persona_heading.paragraph_format.space_after = Pt(3)
        persona_run = persona_heading.add_run(f"Persona {idx}: {persona_name}")
        _apply_run_style(persona_run, bold=True, size=13, color=dark_blue)

        persona_meta = document.add_paragraph()
        persona_meta.paragraph_format.space_after = Pt(10)
        meta_run = persona_meta.add_run(_persona_summary_line(persona))
        _apply_run_style(meta_run, size=10)

        current_section: Optional[str] = None
        for qa in qa_data:
            question = _transcript_plain_text(qa.get("question"))
            answer = _transcript_plain_text(qa.get("answer"))
            metadata = qa.get("metadata") or {}
            section_name = metadata.get("section") or "Discussion Guide"
            answer_meta = generated_answers.get(qa.get("question"), {}) if isinstance(generated_answers, dict) else {}

            if current_section != section_name:
                current_section = section_name
                section_para = document.add_paragraph()
                section_para.paragraph_format.space_before = Pt(8)
                section_para.paragraph_format.space_after = Pt(4)
                section_run = section_para.add_run(str(section_name))
                _apply_run_style(section_run, bold=True, size=11, color=blue)

            question_para = document.add_paragraph(style="List Paragraph")
            question_para.paragraph_format.space_after = Pt(2)
            question_run = question_para.add_run(question)
            _apply_run_style(question_run, size=10)

            answer_para = document.add_paragraph()
            answer_para.paragraph_format.space_after = Pt(2)
            answer_run = answer_para.add_run(f"Respondent: {answer}")
            _apply_run_style(answer_run, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)

            _add_metric_paragraph(
                document,
                answer_meta.get("quality_score"),
                answer_meta.get("independence_score"),
            )

        followups = _extract_follow_up_pairs(interview.get("messages", []))
        if followups:
            followup_heading = document.add_paragraph()
            followup_heading.paragraph_format.space_before = Pt(8)
            followup_heading.paragraph_format.space_after = Pt(4)
            followup_run = followup_heading.add_run("Follow-up Discussion")
            _apply_run_style(followup_run, bold=True, size=11, color=blue)

            for followup in followups:
                moderator_para = document.add_paragraph()
                moderator_para.paragraph_format.space_after = Pt(2)
                moderator_label = moderator_para.add_run("Moderator: ")
                _apply_run_style(moderator_label, bold=True, size=10)
                moderator_text = moderator_para.add_run(_transcript_plain_text(followup.get("question")))
                _apply_run_style(moderator_text, size=10)

                response_para = document.add_paragraph()
                response_para.paragraph_format.space_after = Pt(6)
                response_run = response_para.add_run(
                    f"Respondent: {_transcript_plain_text(followup.get('answer'))}"
                )
                _apply_run_style(response_run, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    document.save(out_path)
    return out_path


async def generate_qual_transcripts_docx(
    objective_id: str,
    out_path: str,
    interview_id: Optional[str] = None,
) -> str:
    interviews = await get_interviews_by_exploration_id(objective_id)

    if interview_id:
        interviews = [
            row for row in interviews
            if row.get("interview_id") == interview_id or row.get("id") == interview_id
        ]

    if not interviews:
        raise ValueError("No interviews found for transcript export")

    research_objective = await get_description(objective_id)
    prepared_interviews: List[Dict[str, Any]] = []
    for interview in interviews:
        qa_data = extract_interview_qa(interview.get("messages", []))
        if not qa_data:
            continue

        persona = await get_persona(interview.get("persona_id")) or {}
        prepared = dict(interview)
        prepared["_persona"] = persona
        prepared_interviews.append(prepared)

    if not prepared_interviews:
        raise ValueError("No valid interview transcript data found")

    return await asyncio.to_thread(
        _build_qual_transcripts_docx,
        objective_id=objective_id,
        research_objective=research_objective,
        interviews=prepared_interviews,
        out_path=out_path,
    )


async def generate_qual_transcripts_csv(objective_id: str) -> bytes:
    """
    Builds a verbatim Q&A CSV from all interviews for the exploration.
    No LLM call — pure data extraction.
    Columns: Persona, Question, Answer, Quality Score, Independence Score
    """
    interviews = await get_interviews_by_exploration_id(objective_id)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Persona", "Question", "Answer", "Quality Score", "Independence Score"])

    for interview in interviews:
        persona_id = interview.get("persona_id")
        persona_details = await get_persona_details(persona_id) or {}
        persona_name = persona_details.get("name", "Unknown")
        qa_data = extract_interview_qa(interview.get("messages", []))

        for qa in qa_data:
            meta = qa.get("metadata", {})
            writer.writerow([
                persona_name,
                qa.get("question", ""),
                qa.get("answer", ""),
                meta.get("quality_score", ""),
                meta.get("independence_score", ""),
            ])

    return output.getvalue().encode("utf-8-sig")  # BOM for Excel compatibility
